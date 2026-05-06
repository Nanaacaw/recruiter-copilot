import os
import re
import fitz
from docx import Document


EMAIL_PATTERN = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_PATTERN = re.compile(
    r"(?<!\w)(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,5}(?!\w)"
)

NAME_LABEL_PATTERN = re.compile(r"^(?:name|nama|full\s*name|candidate)\s*[:\-]\s*(.+)$", re.IGNORECASE)
CONTACT_HINT_PATTERN = re.compile(
    r"(email|e-mail|mail|phone|mobile|tel|telepon|hp|whatsapp|wa|address|alamat|linkedin|github|portfolio|website|www\.|http)",
    re.IGNORECASE,
)
ROLE_HINT_PATTERN = re.compile(
    r"(engineer|developer|designer|manager|specialist|analyst|consultant|scientist|student|fresh graduate|curriculum vitae|resume|cv\b|profile|summary|experience|education)",
    re.IGNORECASE,
)
LOCATION_HINT_PATTERN = re.compile(
    r"\b(indonesia|jakarta|bandung|surabaya|yogyakarta|jogja|bekasi|depok|tangerang|lampung|sumatra|sumatera|java|jawa|bali|regency|city|kota|kabupaten|province|provinsi)\b",
    re.IGNORECASE,
)
NAME_WORD_PATTERN = re.compile(r"^[A-Za-zÀ-ÖØ-öø-ÿ'. -]+$")


class CVParserService:
    def parse(self, file_path: str) -> dict:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            text = self._parse_pdf(file_path)
        elif ext == ".docx":
            text = self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        return self._extract_structured_data(text)

    def _parse_pdf(self, file_path: str) -> str:
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return self._normalize_whitespace(text)

    def _parse_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return self._normalize_whitespace(text)

    def _normalize_whitespace(self, text: str) -> str:
        """Collapse 3+ consecutive blank lines to a single blank line."""
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _extract_structured_data(self, text: str) -> dict:
        lines = self._normalize_lines(text)
        email = self._extract_email(text)
        phone = self._extract_phone(lines, email)
        name = self._extract_name(lines, email, phone)

        return {
            "raw_text": text,
            "name": name,
            "email": email,
            "phone": phone,
            "sections": lines,
            "parse_confidence": {
                "name": "high" if name != "Unknown" else "low",
                "email": "high" if email else "low",
                "phone": "medium" if phone else "low",
            },
        }

    def _normalize_lines(self, text: str) -> list[str]:
        return [
            re.sub(r"\s+", " ", line).strip(" \t\r\n|•·")
            for line in text.splitlines()
            if re.sub(r"\s+", " ", line).strip(" \t\r\n|•·")
        ]

    def _extract_email(self, text: str) -> str:
        match = EMAIL_PATTERN.search(text)
        return match.group(0).strip(".,;:()[]{}<>") if match else ""

    def _extract_phone(self, lines: list[str], email: str) -> str:
        searchable_lines = lines[:35] if lines else []

        for line in searchable_lines:
            clean_line = line.replace(email, " ") if email else line
            clean_line = re.sub(r"[–—−]", "-", clean_line)
            clean_line = re.sub(r"(?<=\d)\s*\?\s*(?=\d)", "-", clean_line)
            clean_line = re.sub(r"(?<=\d)\s*-\s*(?=\d)", "-", clean_line)
            clean_line = re.sub(r"(?<=\d)\s*\.\s*(?=\d)", "-", clean_line)

            clean_line = re.sub(
                r"\S*(?:linkedin|github|http|www\.|portfolio)\S*",
                " ",
                clean_line,
                flags=re.IGNORECASE,
            )

            for match in PHONE_PATTERN.finditer(clean_line):
                candidate = match.group(0).strip(".,;:()[]{}<> ")
                digits = re.sub(r"\D", "", candidate)

                if 9 <= len(digits) <= 15 and not self._looks_like_date_or_year(digits):
                    return self._normalize_phone(candidate)

        return ""

    def _extract_name(self, lines: list[str], email: str, phone: str) -> str:
        for line in lines[:12]:
            labeled_name = NAME_LABEL_PATTERN.match(line)
            if labeled_name:
                candidate = self._clean_name(labeled_name.group(1))
                if self._is_probable_name(candidate):
                    return candidate

        candidates: list[tuple[int, str]] = []
        for index, line in enumerate(lines[:18]):
            candidate = self._clean_name(line)
            if not candidate or candidate == email or candidate == phone:
                continue
            if not self._is_probable_name(candidate):
                continue

            score = 100 - index * 5
            if index <= 3:
                score += 20
            if candidate.isupper():
                score += 8
            if 2 <= len(candidate.split()) <= 4:
                score += 10

            candidates.append((score, candidate.title() if candidate.isupper() else candidate))

        if not candidates:
            return "Unknown"

        candidates.sort(key=lambda item: item[0], reverse=True)
        return candidates[0][1]

    def _clean_name(self, value: str) -> str:
        cleaned = EMAIL_PATTERN.sub(" ", value)
        cleaned = PHONE_PATTERN.sub(" ", cleaned)
        cleaned = re.sub(r"[\|•·_/\\]+", " ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned)
        return cleaned.strip(" -:,.")

    def _is_probable_name(self, value: str) -> bool:
        if not value:
            return False
        if CONTACT_HINT_PATTERN.search(value) or ROLE_HINT_PATTERN.search(value):
            return False
        if LOCATION_HINT_PATTERN.search(value) or "," in value:
            return False
        if any(char.isdigit() for char in value):
            return False
        if "@" in value or len(value) > 70:
            return False

        words = [word.strip(".,") for word in value.split() if word.strip(".,")]
        if not 2 <= len(words) <= 5:
            return False
        if not NAME_WORD_PATTERN.match(value):
            return False
        if sum(len(word) <= 1 for word in words) > 1:
            return False

        return True

    def _looks_like_link_or_id_line(self, value: str) -> bool:
        normalized = value.lower()
        return (
            "linkedin" in normalized
            or "github" in normalized
            or "http" in normalized
            or "www." in normalized
            or "portfolio" in normalized
        )

    def _looks_like_date_or_year(self, digits: str) -> bool:
        return len(digits) in {4, 6, 8} and digits.startswith(("19", "20"))

    def _normalize_phone(self, value: str) -> str:
        value = re.sub(r"\s+", " ", value).strip()
        value = re.sub(r"(?<=\d)\s*\?\s*(?=\d)", "-", value)
        value = re.sub(r"(?<=\d)\s*-\s*(?=\d)", "-", value)
        value = re.sub(r"(?<=\d)\s*\.\s*(?=\d)", "-", value)
        return value


cv_parser_service = CVParserService()
